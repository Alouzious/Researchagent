from typing import List, Dict, Optional
from datetime import datetime
from agents.llm_agent import llm_agent
from agents.export_agent import export_agent
from utils.error_handler import logger, handle_errors, ErrorContext
from database.database import db
from database.models import LiteratureReview


class LiteratureReviewAgent:
    """Orchestrates literature review generation"""
    
    def __init__(self):
        self.llm_agent = llm_agent
        self.export_agent = export_agent
        logger.info("Literature Review Agent initialized")
    
    @handle_errors#(default_return=None)
    def generate_review(
        self,
        papers: List[Dict],
        query: str,
        detail_level: str = "medium",
        review_type: str = "thematic",
        project_id: Optional[int] = None,
        save_to_db: bool = True
    ) -> Dict:
        """
        Generate a complete literature review
        
        Args:
            papers: List of paper dictionaries (must include summaries)
            query: Research topic/query
            detail_level: 'short', 'medium', or 'long'
            review_type: 'thematic', 'chronological', or 'methodological'
            project_id: Optional project ID to associate with
            save_to_db: Whether to save to database
            
        Returns:
            Dictionary containing review content and metadata
        """
        with ErrorContext(f"Generating literature review for: {query}"):
            logger.info(f"Starting literature review generation: {len(papers)} papers, {detail_level} detail")
            
            # Validate papers have required data
            valid_papers = self._validate_papers(papers)
            if len(valid_papers) < 3:
                raise ValueError("Need at least 3 papers with summaries to generate a review")
            
            # Generate the review
            review_content = self.llm_agent.generate_literature_review(
                papers=valid_papers,
                query=query,
                detail_level=detail_level,
                review_type=review_type
            )
            
            # Generate export formats
            markdown_content = review_content  # Already in markdown
            latex_content = self._convert_to_latex(review_content, query)
            
            # Prepare result
            result = {
                'title': f"Literature Review: {query}",
                'content': review_content,
                'markdown_content': markdown_content,
                'latex_content': latex_content,
                'query': query,
                'detail_level': detail_level,
                'review_type': review_type,
                'papers_count': len(valid_papers),
                'paper_ids': [p.get('paper_id') or p.get('title') for p in valid_papers],
                'generated_at': datetime.now().isoformat(),
                'word_count': len(review_content.split()),
                'page_estimate': len(review_content.split()) // 250  # ~250 words per page
            }
            
            # Save to database if requested
            if save_to_db:
                try:
                    lit_review = LiteratureReview(
                        project_id=project_id,
                        title=result['title'],
                        content=review_content,
                        papers_included=[p.get('title') for p in valid_papers],
                        review_type=review_type,
                        detail_level=detail_level,
                        markdown_content=markdown_content,
                        latex_content=latex_content
                    )
                    
                    with db.get_session() as session:
                        session.add(lit_review)
                        session.flush()
                        result['review_id'] = lit_review.id
                    
                    logger.info(f"Literature review saved to database with ID: {lit_review.id}")
                except Exception as e:
                    logger.warning(f"Failed to save literature review to database: {str(e)}")
            
            logger.info(f"Literature review generated: {result['word_count']} words, ~{result['page_estimate']} pages")
            return result
    
    def _validate_papers(self, papers: List[Dict]) -> List[Dict]:
        """Validate papers have required data for review generation"""
        valid_papers = []
        
        for paper in papers:
            # Must have at least title and abstract or summary
            if paper.get('title') and (paper.get('abstract') or paper.get('abstract_summary')):
                valid_papers.append(paper)
        
        logger.info(f"Validated {len(valid_papers)}/{len(papers)} papers for review generation")
        return valid_papers
    
    def _convert_to_latex(self, markdown_content: str, title: str) -> str:
        """Convert markdown review to LaTeX"""
        # Start with LaTeX document
        latex_lines = [
            r"\documentclass[11pt]{article}",
            r"\usepackage[utf8]{inputenc}",
            r"\usepackage{hyperref}",
            r"\usepackage{geometry}",
            r"\geometry{margin=1in}",
            r"\usepackage{parskip}",
            f"\\title{{{self._escape_latex(title)}}}",
            r"\author{AI Research Agent}",
            r"\date{\today}",
            r"\begin{document}",
            r"\maketitle",
            ""
        ]
        
        # Convert markdown to LaTeX
        lines = markdown_content.split('\n')
        in_code_block = False
        
        for line in lines:
            # Skip title (already in \title)
            if line.startswith('# '):
                continue
            
            # Section headers
            if line.startswith('## '):
                section_title = self._escape_latex(line[3:])
                latex_lines.append(f"\\section{{{section_title}}}")
            elif line.startswith('### '):
                subsection_title = self._escape_latex(line[4:])
                latex_lines.append(f"\\subsection{{{subsection_title}}}")
            elif line.startswith('#### '):
                subsubsection_title = self._escape_latex(line[5:])
                latex_lines.append(f"\\subsubsection{{{subsubsection_title}}}")
            
            # Code blocks
            elif line.startswith('```'):
                if in_code_block:
                    latex_lines.append(r"\end{verbatim}")
                    in_code_block = False
                else:
                    latex_lines.append(r"\begin{verbatim}")
                    in_code_block = True
            
            # Regular text
            elif line.strip():
                if not in_code_block:
                    # Handle bold and italic
                    text = self._escape_latex(line)
                    text = text.replace('**', r'\textbf{').replace('**', '}')
                    text = text.replace('*', r'\textit{').replace('*', '}')
                    latex_lines.append(text)
                else:
                    latex_lines.append(line)
            
            # Empty lines
            else:
                latex_lines.append("")
        
        latex_lines.extend([
            "",
            r"\end{document}"
        ])
        
        return "\n".join(latex_lines)
    
    def _escape_latex(self, text: str) -> str:
        """Escape special LaTeX characters"""
        if not text:
            return ""
        
        replacements = {
            '&': r'\&',
            '%': r'\%',
            '$': r'\$',
            '#': r'\#',
            '_': r'\_',
            '{': r'\{',
            '}': r'\}',
            '~': r'\textasciitilde{}',
            '^': r'\^{}',
            '\\': r'\textbackslash{}',
        }
        
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)
        
        return text
    
    @handle_errors#(default_return=None)
    def export_review(
        self,
        review_content: str,
        format: str,
        filename: str = None
    ):
        """
        Export literature review to specified format
        
        Args:
            review_content: Review text (markdown)
            format: Export format ('word', 'pdf', 'latex', 'markdown')
            filename: Optional filename
            
        Returns:
            Exported content (BytesIO or string)
        """
        if format == 'markdown':
            return review_content
        
        elif format == 'latex':
            return self._convert_to_latex(review_content, "Literature Review")
        
        elif format == 'word':
            # Convert to structured data for Word export
            sections = self._parse_markdown_sections(review_content)
            return self._export_to_word(sections, filename)
        
        elif format == 'pdf':
            sections = self._parse_markdown_sections(review_content)
            return self._export_to_pdf(sections, filename)
        
        else:
            raise ValueError(f"Unsupported export format: {format}")
    
    def _parse_markdown_sections(self, markdown: str) -> Dict:
        """Parse markdown into structured sections"""
        sections = {
            'title': '',
            'sections': []
        }
        
        lines = markdown.split('\n')
        current_section = None
        
        for line in lines:
            if line.startswith('# '):
                sections['title'] = line[2:].strip()
            elif line.startswith('## '):
                if current_section:
                    sections['sections'].append(current_section)
                current_section = {
                    'title': line[3:].strip(),
                    'content': []
                }
            elif current_section and line.strip():
                current_section['content'].append(line)
        
        if current_section:
            sections['sections'].append(current_section)
        
        return sections
    
    def _export_to_word(self, sections: Dict, filename: str = None):
        """Export structured review to Word"""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        
        doc = Document()
        
        # Title
        title = doc.add_heading(sections.get('title', 'Literature Review'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph()
        
        # Sections
        for section in sections.get('sections', []):
            doc.add_heading(section['title'], level=1)
            content_text = '\n'.join(section['content'])
            doc.add_paragraph(content_text)
            doc.add_paragraph()
        
        # Save to BytesIO
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes
    
    def _export_to_pdf(self, sections: Dict, filename: str = None):
        """Export structured review to PDF"""
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        import io
        
        pdf_bytes = io.BytesIO()
        doc = SimpleDocTemplate(pdf_bytes, pagesize=letter)
        styles = getSampleStyleSheet()
        
        story = []
        
        # Title
        story.append(Paragraph(sections.get('title', 'Literature Review'), styles['Title']))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Spacer(1, 0.5*inch))
        
        # Sections
        for section in sections.get('sections', []):
            story.append(Paragraph(section['title'], styles['Heading1']))
            story.append(Spacer(1, 0.2*inch))
            
            content_text = '\n'.join(section['content'])
            story.append(Paragraph(content_text, styles['Normal']))
            story.append(Spacer(1, 0.3*inch))
        
        doc.build(story)
        pdf_bytes.seek(0)
        
        return pdf_bytes
    
    def get_review_statistics(self, review_content: str) -> Dict:
        """Get statistics about the generated review"""
        words = review_content.split()
        
        # Count sections
        sections = review_content.count('##')
        
        # Estimate reading time (200 words per minute)
        reading_time = len(words) // 200
        
        return {
            'word_count': len(words),
            'character_count': len(review_content),
            'section_count': sections,
            'estimated_pages': len(words) // 250,
            'reading_time_minutes': reading_time,
            'references_count': review_content.count('\n\n')  # Rough estimate
        }


# Global literature review agent instance
lit_review_agent = LiteratureReviewAgent()
