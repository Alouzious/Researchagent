
import io
import json
import csv
from typing import List, Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_JUSTIFY
import bibtexparser
from bibtexparser.bwriter import BibTexWriter
from utils.error_handler import logger, handle_errors


class ExportAgent:
    """Handles exporting research data to various formats"""
    
    def __init__(self):
        logger.info("Export Agent initialized")
    
    # ==================== WORD (DOCX) EXPORT ====================
    
    @handle_errors#(default_return=None)
    def export_to_word(self, papers: List[Dict], filename: str = None, include_summaries: bool = True) -> io.BytesIO:
        """
        Export papers to Word document
        
        Args:
            papers: List of paper dictionaries
            filename: Optional filename
            include_summaries: Include AI summaries
            
        Returns:
            BytesIO object containing the Word document
        """
        doc = Document()
        
        # Title
        title = doc.add_heading('Research Papers Summary', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
        doc.add_paragraph(f'Total Papers: {len(papers)}')
        doc.add_paragraph('')
        
        # Papers
        for i, paper in enumerate(papers, 1):
            # Paper title
            doc.add_heading(f'{i}. {paper.get("title", "Untitled")}', level=2)
            
            # Metadata
            authors = ", ".join(paper.get('authors', [])[:5])
            if len(paper.get('authors', [])) > 5:
                authors += " et al."
            
            metadata = doc.add_paragraph()
            metadata.add_run(f'Authors: ').bold = True
            metadata.add_run(f'{authors}\n')
            metadata.add_run(f'Year: ').bold = True
            metadata.add_run(f'{paper.get("year", "Unknown")}\n')
            metadata.add_run(f'Venue: ').bold = True
            metadata.add_run(f'{paper.get("venue", "Unknown")}\n')
            metadata.add_run(f'Citations: ').bold = True
            metadata.add_run(f'{paper.get("citation_count", 0)}\n')
            metadata.add_run(f'URL: ').bold = True
            metadata.add_run(f'{paper.get("url", "N/A")}')
            
            # Abstract
            if paper.get('abstract'):
                doc.add_heading('Abstract', level=3)
                doc.add_paragraph(paper['abstract'])
            
            # Summaries
            if include_summaries:
                if paper.get('abstract_summary'):
                    doc.add_heading('Summary', level=3)
                    doc.add_paragraph(paper['abstract_summary'])
                
                if paper.get('key_findings'):
                    doc.add_heading('Key Findings', level=3)
                    doc.add_paragraph(paper['key_findings'])
                
                # Research Gaps
                if paper.get('research_gaps'):
                    gaps = paper['research_gaps']
                    
                    doc.add_heading('Research Gaps', level=3)
                    
                    if gaps.get('methodology_gaps'):
                        doc.add_paragraph('Methodology Gaps:', style='Heading 4')
                        doc.add_paragraph(gaps['methodology_gaps'])
                    
                    if gaps.get('knowledge_gaps'):
                        doc.add_paragraph('Knowledge Gaps:', style='Heading 4')
                        doc.add_paragraph(gaps['knowledge_gaps'])
                    
                    if gaps.get('future_directions'):
                        doc.add_paragraph('Future Directions:', style='Heading 4')
                        doc.add_paragraph(gaps['future_directions'])
            
            # Page break between papers (except last)
            if i < len(papers):
                doc.add_page_break()
        
        # Save to BytesIO
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        logger.info(f"Exported {len(papers)} papers to Word")
        return doc_bytes
    
    # ==================== PDF EXPORT ====================
    
    @handle_errors#(default_return=None)
    def export_to_pdf(self, papers: List[Dict], filename: str = None) -> io.BytesIO:
        """Export papers to PDF"""
        pdf_bytes = io.BytesIO()
        doc = SimpleDocTemplate(pdf_bytes, pagesize=letter)
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=RGBColor(0, 0, 128),
            spaceAfter=30,
            alignment=TA_JUSTIFY
        )
        
        # Build PDF content
        story = []
        
        # Title
        story.append(Paragraph("Research Papers Summary", title_style))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Paragraph(f"Total Papers: {len(papers)}", styles['Normal']))
        story.append(Spacer(1, 0.5*inch))
        
        # Papers
        for i, paper in enumerate(papers, 1):
            # Title
            story.append(Paragraph(f"{i}. {paper.get('title', 'Untitled')}", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            
            # Metadata
            authors = ", ".join(paper.get('authors', [])[:5])
            if len(paper.get('authors', [])) > 5:
                authors += " et al."
            
            story.append(Paragraph(f"<b>Authors:</b> {authors}", styles['Normal']))
            story.append(Paragraph(f"<b>Year:</b> {paper.get('year', 'Unknown')}", styles['Normal']))
            story.append(Paragraph(f"<b>Citations:</b> {paper.get('citation_count', 0)}", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
            
            # Abstract summary
            if paper.get('abstract_summary'):
                story.append(Paragraph("<b>Summary:</b>", styles['Heading3']))
                story.append(Paragraph(paper['abstract_summary'], styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
            
            if i < len(papers):
                story.append(PageBreak())
        
        doc.build(story)
        pdf_bytes.seek(0)
        
        logger.info(f"Exported {len(papers)} papers to PDF")
        return pdf_bytes
    
    # ==================== CSV EXPORT ====================
    
    @handle_errors#(default_return=None)
    def export_to_csv(self, papers: List[Dict]) -> io.StringIO:
        """Export papers to CSV"""
        csv_buffer = io.StringIO()
        
        if not papers:
            return csv_buffer
        
        # Define columns
        fieldnames = [
            'title', 'authors', 'year', 'venue', 'citation_count',
            'doi', 'url', 'abstract', 'summary', 'has_pdf'
        ]
        
        writer = csv.DictWriter(csv_buffer, fieldnames=fieldnames)
        writer.writeheader()
        
        for paper in papers:
            row = {
                'title': paper.get('title', ''),
                'authors': '; '.join(paper.get('authors', [])),
                'year': paper.get('year', ''),
                'venue': paper.get('venue', ''),
                'citation_count': paper.get('citation_count', 0),
                'doi': paper.get('doi', ''),
                'url': paper.get('url', ''),
                'abstract': paper.get('abstract', ''),
                'summary': paper.get('abstract_summary', ''),
                'has_pdf': paper.get('has_pdf', False)
            }
            writer.writerow(row)
        
        csv_buffer.seek(0)
        logger.info(f"Exported {len(papers)} papers to CSV")
        return csv_buffer
    
    # ==================== BIBTEX EXPORT ====================
    
    @handle_errors#(default_return=None)
    def export_to_bibtex(self, papers: List[Dict]) -> str:
        """Export papers to BibTeX format"""
        db = bibtexparser.bibdatabase.BibDatabase()
        
        for i, paper in enumerate(papers, 1):
            # Generate citation key
            first_author = paper.get('authors', ['unknown'])[0].split()[-1].lower()
            year = paper.get('year', 'nodate')
            key = f"{first_author}{year}_{i}"
            
            entry = {
                'ENTRYTYPE': 'article',
                'ID': key,
                'title': paper.get('title', 'Untitled'),
                'author': ' and '.join(paper.get('authors', ['Unknown'])),
                'year': str(paper.get('year', '')),
                'journal': paper.get('venue', ''),
                'url': paper.get('url', ''),
                'abstract': paper.get('abstract', '')
            }
            
            if paper.get('doi'):
                entry['doi'] = paper['doi']
            
            db.entries.append(entry)
        
        writer = BibTexWriter()
        bibtex_str = writer.write(db)
        
        logger.info(f"Exported {len(papers)} papers to BibTeX")
        return bibtex_str
    
    # ==================== MARKDOWN EXPORT ====================
    
    @handle_errors#(default_return=None)
    def export_to_markdown(self, papers: List[Dict], include_summaries: bool = True) -> str:
        """Export papers to Markdown format"""
        md_lines = []
        
        # Header
        md_lines.append("# Research Papers Summary\n")
        md_lines.append(f"Generated on: {datetime.now().strftime('%B %d, %Y')}\n")
        md_lines.append(f"Total Papers: {len(papers)}\n")
        md_lines.append("---\n\n")
        
        # Papers
        for i, paper in enumerate(papers, 1):
            md_lines.append(f"## {i}. {paper.get('title', 'Untitled')}\n\n")
            
            # Metadata
            authors = ", ".join(paper.get('authors', [])[:5])
            if len(paper.get('authors', [])) > 5:
                authors += " et al."
            
            md_lines.append(f"**Authors:** {authors}  \n")
            md_lines.append(f"**Year:** {paper.get('year', 'Unknown')}  \n")
            md_lines.append(f"**Venue:** {paper.get('venue', 'Unknown')}  \n")
            md_lines.append(f"**Citations:** {paper.get('citation_count', 0)}  \n")
            md_lines.append(f"**URL:** {paper.get('url', 'N/A')}  \n\n")
            
            # Abstract
            if paper.get('abstract'):
                md_lines.append("### Abstract\n\n")
                md_lines.append(f"{paper['abstract']}\n\n")
            
            # Summary
            if include_summaries and paper.get('abstract_summary'):
                md_lines.append("### Summary\n\n")
                md_lines.append(f"{paper['abstract_summary']}\n\n")
            
            # Key Findings
            if include_summaries and paper.get('key_findings'):
                md_lines.append("### Key Findings\n\n")
                md_lines.append(f"{paper['key_findings']}\n\n")
            
            # Research Gaps
            if include_summaries and paper.get('research_gaps'):
                gaps = paper['research_gaps']
                md_lines.append("### Research Gaps\n\n")
                
                if gaps.get('methodology_gaps'):
                    md_lines.append("#### Methodology Gaps\n\n")
                    md_lines.append(f"{gaps['methodology_gaps']}\n\n")
                
                if gaps.get('knowledge_gaps'):
                    md_lines.append("#### Knowledge Gaps\n\n")
                    md_lines.append(f"{gaps['knowledge_gaps']}\n\n")
                
                if gaps.get('future_directions'):
                    md_lines.append("#### Future Directions\n\n")
                    md_lines.append(f"{gaps['future_directions']}\n\n")
            
            md_lines.append("---\n\n")
        
        markdown = "".join(md_lines)
        logger.info(f"Exported {len(papers)} papers to Markdown")
        return markdown
    
    # ==================== LATEX EXPORT ====================
    
    @handle_errors#(default_return=None)
    def export_to_latex(self, papers: List[Dict]) -> str:
        """Export papers to LaTeX format"""
        latex_lines = []
        
        # Document header
        latex_lines.append(r"\documentclass[11pt]{article}")
        latex_lines.append(r"\usepackage[utf8]{inputenc}")
        latex_lines.append(r"\usepackage{hyperref}")
        latex_lines.append(r"\usepackage{geometry}")
        latex_lines.append(r"\geometry{margin=1in}")
        latex_lines.append(r"\title{Research Papers Summary}")
        latex_lines.append(r"\author{AI Research Agent}")
        latex_lines.append(r"\date{\today}")
        latex_lines.append(r"\begin{document}")
        latex_lines.append(r"\maketitle")
        latex_lines.append("")
        
        # Papers
        for i, paper in enumerate(papers, 1):
            # Escape LaTeX special characters
            title = self._escape_latex(paper.get('title', 'Untitled'))
            latex_lines.append(f"\\section{{{title}}}")
            latex_lines.append("")
            
            # Metadata
            authors = ", ".join(paper.get('authors', [])[:5])
            if len(paper.get('authors', [])) > 5:
                authors += " et al."
            authors = self._escape_latex(authors)
            
            latex_lines.append(f"\\textbf{{Authors:}} {authors} \\\\")
            latex_lines.append(f"\\textbf{{Year:}} {paper.get('year', 'Unknown')} \\\\")
            latex_lines.append(f"\\textbf{{Citations:}} {paper.get('citation_count', 0)} \\\\")
            
            if paper.get('url'):
                latex_lines.append(f"\\textbf{{URL:}} \\url{{{paper['url']}}} \\\\")
            
            latex_lines.append("")
            
            # Abstract
            if paper.get('abstract'):
                abstract = self._escape_latex(paper['abstract'])
                latex_lines.append(r"\subsection{Abstract}")
                latex_lines.append(abstract)
                latex_lines.append("")
            
            # Summary
            if paper.get('abstract_summary'):
                summary = self._escape_latex(paper['abstract_summary'])
                latex_lines.append(r"\subsection{Summary}")
                latex_lines.append(summary)
                latex_lines.append("")
            
            latex_lines.append(r"\newpage")
            latex_lines.append("")
        
        latex_lines.append(r"\end{document}")
        
        latex_content = "\n".join(latex_lines)
        logger.info(f"Exported {len(papers)} papers to LaTeX")
        return latex_content
    
    def _escape_latex(self, text: str) -> str:
        """Escape special LaTeX characters"""
        if not text:
            return ""
        
        replacements = {
            '&': r'\&',
            '%': r'\%',
            '$': r'\$',
            '#': r'\#',
            '_': r'\_',
            '{': r'\{',
            '}': r'\}',
            '~': r'\textasciitilde{}',
            '^': r'\^{}',
            '\\': r'\textbackslash{}',
        }
        
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)
        
        return text
    
    # ==================== JSON EXPORT ====================
    
    @handle_errors#(default_return=None)
    def export_to_json(self, papers: List[Dict]) -> str:
        """Export papers to JSON format"""
        json_data = {
            'exported_at': datetime.now().isoformat(),
            'total_papers': len(papers),
            'papers': papers
        }
        
        json_str = json.dumps(json_data, indent=2, ensure_ascii=False)
        logger.info(f"Exported {len(papers)} papers to JSON")
        return json_str


# Global export agent instance
export_agent = ExportAgent()